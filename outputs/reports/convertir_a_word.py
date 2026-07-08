"""Convertir informe academico de Markdown a Word (.docx) con formato profesional."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Configuracion de estilos ─────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1A, 0x52, 0x76)  # Azul oscuro
COLOR_ACCENT = RGBColor(0x29, 0x80, 0xB9)   # Azul medio
COLOR_DARK = RGBColor(0x2C, 0x3E, 0x50)     # Texto oscuro
COLOR_GRAY = RGBColor(0x7F, 0x8C, 0x8D)     # Texto gris
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_HEADER = RGBColor(0x1A, 0x52, 0x76)

FONT_TITLE = "Calibri"
FONT_BODY = "Calibri"

# Ruta de figuras
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
FIGURES_DIR = PROJECT_ROOT / "outputs" / "figures"

def set_cell_shading(cell, color):
    """Aplicar color de fondo a una celda."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_page_number(doc):
    """Agregar numero de pagina al footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

def add_cover_page(doc):
    """Crear portada profesional."""
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()

    # Titulo principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME DE PROYECTO")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True
    run.font.name = FONT_TITLE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DE ANÁLISIS DE DATOS")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.bold = True
    run.font.name = FONT_TITLE

    doc.add_paragraph()

    # Linea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Subtitulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Modelado Predictivo de la Severidad de\n"
        "Siniestros Viales Fatales en el Perú\n"
        "(ONSV 2021–2025)"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_DARK
    run.font.name = FONT_TITLE

    for _ in range(3):
        doc.add_paragraph()

    # Informacion institucional
    info_lines = [
        "Universidad Nacional del Altiplano",
        "Escuela Profesional de Ingeniería de Sistemas",
        "Julio 2025",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_GRAY
        run.font.name = FONT_TITLE

    # Salto de pagina
    doc.add_page_break()


def create_styles(doc):
    """Crear estilos personalizados para el documento."""
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def parse_markdown_to_docx(md_path, doc):
    """Parsear archivo Markdown y agregar contenido al documento."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    table_mode = False
    table_lines = []
    code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Saltar lineas de separacion decorativas
        if line.startswith("---") and len(line) > 3:
            i += 1
            continue

        # Saltar el titulo principal ya que esta en la portada
        if line.startswith("# INFORME"):
            i += 1
            continue
        if line.startswith("**Título:**") or line.startswith("**Autores:**") or \
           line.startswith("**Institución:**") or line.startswith("**Fecha:**"):
            i += 1
            continue

        # Detectar inicio/fin de bloque de codigo
        if line.startswith("```"):
            code_block = not code_block
            i += 1
            continue

        if code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_DARK
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1)
            i += 1
            continue

        # Detectar tablas
        if "|" in line and line.strip().startswith("|"):
            table_lines.append(line)
            table_mode = True
            i += 1
            continue
        else:
            if table_mode and len(table_lines) > 0:
                process_table(table_lines, doc)
                table_lines = []
                table_mode = False

        # Headers
        if line.startswith("## "):
            title = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = FONT_TITLE
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif line.startswith("### "):
            title = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = COLOR_ACCENT
            run.font.name = FONT_TITLE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)

        elif line.startswith("#### "):
            title = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = COLOR_DARK
            p.paragraph_format.space_before = Pt(10)

        elif line.startswith("1. ") or line.startswith("2. ") or \
             line.startswith("3. ") or line.startswith("4. ") or \
             line.startswith("5. ") or line.startswith("6. ") or \
             line.startswith("7. ") or line.startswith("8. ") or \
             line.startswith("9. ") or line.startswith("10."):
            # Lista numerada
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph()
            run = p.add_run(f"    {text}")
            run.font.size = Pt(11)
            run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(2)

        elif line.startswith("- "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(f"    • {text}")
            run.font.size = Pt(11)
            run.font.name = FONT_BODY
            p.paragraph_format.space_before = Pt(2)

        # Detectar imagenes ![alt](ruta)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            match = re.match(r'^\!\[(.*?)\]\((.*?)\)$', line)
            if match:
                alt_text = match.group(1)
                img_rel = match.group(2)
                # Buscar imagen en outputs/figures/
                img_path = FIGURES_DIR / img_rel
                if img_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    # Pie de imagen
                    p2 = doc.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run2 = p2.add_run(f"Figura: {alt_text}")
                    run2.font.size = Pt(9)
                    run2.font.italic = True
                    run2.font.color.rgb = COLOR_GRAY
                    run2.font.name = FONT_BODY

        elif line.strip() == "":
            # Parrafo vacio
            doc.add_paragraph()

        else:
            # Parrafo normal con formato
            text = line.strip()
            if text:
                p = doc.add_paragraph()
                # Procesar negritas **texto**
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.name = FONT_BODY
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        run.font.name = FONT_BODY
                p.paragraph_format.space_before = Pt(3)

        i += 1

    # Procesar tabla final si existe
    if table_mode and len(table_lines) > 0:
        process_table(table_lines, doc)


def process_table(lines, doc):
    """Convertir lineas de tabla Markdown a tabla de Word."""
    if len(lines) < 2:
        return

    # Limpiar y procesar
    clean_lines = []
    for l in lines:
        l = l.strip()
        if l.startswith("|") and l.endswith("|"):
            l = l[1:-1]
        elif l.startswith("|"):
            l = l[1:]
        clean_lines.append(l)

    # Saltar la linea de separacion (|---|---|)
    data_lines = [l for l in clean_lines if not re.match(r'^[\s\-:|]+$', l)]
    if len(data_lines) < 1:
        return

    # Determinar numero de columnas
    header = [c.strip() for c in data_lines[0].split("|")]
    n_cols = len(header)

    # Crear tabla
    table = doc.add_table(rows=len(data_lines), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Llenar datos
    for row_idx, line in enumerate(data_lines):
        cells = [c.strip() for c in line.split("|")]
        for col_idx in range(n_cols):
            cell = table.cell(row_idx, col_idx)
            text = cells[col_idx] if col_idx < len(cells) else ""
            cell.text = ""

            # Formato del texto
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = FONT_BODY

            if row_idx == 0:
                # Fila de encabezado
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
                run.font.size = Pt(9)
                set_cell_shading(cell, "1A5276")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Resaltar filas alternadas
            elif row_idx % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()  # Espacio despues de la tabla


def main():
    """Funcion principal."""
    script_dir = Path(__file__).parent
    md_path = script_dir / "informe_academico.md"
    docx_path = script_dir / "informe_academico.docx"

    if not md_path.exists():
        print(f"Error: No se encuentra {md_path}")
        return

    print("Creando documento Word...")
    doc = Document()

    # Configurar pagina
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Crear portada
    add_cover_page(doc)

    # Crear estilos
    create_styles(doc)

    # Agregar contenido
    print("Agregando contenido...")
    parse_markdown_to_docx(md_path, doc)

    # Agregar numeros de pagina
    add_page_number(doc)

    # Guardar
    doc.save(docx_path)
    print(f"Documento guardado: {docx_path}")
    print(f"Tamanio: {docx_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
